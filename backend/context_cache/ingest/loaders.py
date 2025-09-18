"""Document loaders for supported formats."""

from __future__ import annotations

import email.utils
import mailbox
import os
from datetime import datetime
from email import message_from_binary_file, policy
from email.message import EmailMessage
from pathlib import Path

try:
    import fitz  # type: ignore[attr-defined]
except Exception:  # pragma: no cover - optional dependency
    fitz = None
try:
    import langid
except Exception:  # pragma: no cover - fallback when langid unavailable
    langid = None
try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None
from markdown_it import MarkdownIt
import yaml

from context_cache.ingest.types import LoadedDocument
from context_cache.utils.text import normalize

_MD = MarkdownIt()


class BaseLoader:
    """Common loader interface."""

    suffixes: tuple[str, ...] = ()
    mime_type: str = "application/octet-stream"

    def can_load(self, path: Path) -> bool:
        return path.suffix.lower() in self.suffixes

    def load(self, path: Path) -> list[LoadedDocument]:  # pragma: no cover - interface
        raise NotImplementedError


class MarkdownLoader(BaseLoader):
    suffixes = (".md", ".markdown", ".mdx")
    mime_type = "text/markdown"

    def load(self, path: Path) -> list[LoadedDocument]:
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="ignore")
        front_matter, body = _split_front_matter(text)
        normalized = _markdown_to_text(body)
        metadata = {"path": str(path), "lang": _detect_lang(normalized)}
        if front_matter:
            metadata["front_matter"] = front_matter
        return [
            LoadedDocument(
                path=path,
                text=normalized,
                raw_bytes=raw,
                metadata=metadata,
                mime=self.mime_type,
                title=front_matter.get("title") if front_matter else path.stem,
                author=front_matter.get("author") if front_matter else None,
                created_ts=_to_unix_ms(front_matter.get("created")) if front_matter else None,
                modified_ts=int(path.stat().st_mtime * 1000),
                size_bytes=len(raw),
            )
        ]


class TextLoader(BaseLoader):
    suffixes = (".txt", ".text", ".log")
    mime_type = "text/plain"

    def load(self, path: Path) -> list[LoadedDocument]:
        raw = path.read_bytes()
        text = raw.decode("utf-8", errors="ignore")
        metadata = {"path": str(path), "lang": _detect_lang(text)}
        return [
            LoadedDocument(
                path=path,
                text=normalize(text),
                raw_bytes=raw,
                metadata=metadata,
                mime=self.mime_type,
                title=path.stem,
                author=None,
                created_ts=int(path.stat().st_ctime * 1000),
                modified_ts=int(path.stat().st_mtime * 1000),
                size_bytes=len(raw),
            )
        ]


class PDFLoader(BaseLoader):
    suffixes = (".pdf",)
    mime_type = "application/pdf"

    def load(self, path: Path) -> list[LoadedDocument]:
        if fitz is None:
            raise RuntimeError("PyMuPDF is required to load PDF files")
        raw = path.read_bytes()
        with fitz.open(stream=raw, filetype="pdf") as doc:
            pages = [page.get_text("text", sort=True) for page in doc]
        text = normalize("\n\n".join(pages))
        metadata = {
            "path": str(path),
            "page_count": len(pages),
            "lang": _detect_lang(text),
        }
        return [
            LoadedDocument(
                path=path,
                text=text,
                raw_bytes=raw,
                metadata=metadata,
                mime=self.mime_type,
                title=path.stem,
                author=None,
                created_ts=_file_time(path.stat().st_ctime),
                modified_ts=_file_time(path.stat().st_mtime),
                size_bytes=len(raw),
            )
        ]


class DocxLoader(BaseLoader):
    suffixes = (".docx",)
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def load(self, path: Path) -> list[LoadedDocument]:
        if Document is None:
            raise RuntimeError("python-docx is required to load DOCX files")
        raw = path.read_bytes()
        document = Document(path)
        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
        text = normalize("\n".join(paragraphs))
        core = document.core_properties
        metadata = {
            "path": str(path),
            "lang": _detect_lang(text),
            "category": core.category,
        }
        return [
            LoadedDocument(
                path=path,
                text=text,
                raw_bytes=raw,
                metadata=metadata,
                mime=self.mime_type,
                title=core.title or path.stem,
                author=core.author or None,
                created_ts=_to_unix_ms(core.created),
                modified_ts=_to_unix_ms(core.modified) or _file_time(path.stat().st_mtime),
                size_bytes=len(raw),
            )
        ]


class EmailLoader(BaseLoader):
    suffixes = (".eml",)
    mime_type = "message/rfc822"

    def load(self, path: Path) -> list[LoadedDocument]:
        with path.open("rb") as handle:
            message = message_from_binary_file(handle, policy=policy.default)
        return [_email_to_document(path, message)]


class MboxLoader(BaseLoader):
    suffixes = (".mbox",)
    mime_type = "application/mbox"

    def load(self, path: Path) -> list[LoadedDocument]:
        documents: list[LoadedDocument] = []
        mbox = mailbox.mbox(path, factory=lambda f: message_from_binary_file(f, policy=policy.default))
        for idx, message in enumerate(mbox):
            documents.append(_email_to_document(path, message, index=idx))
        return documents


class LoaderRegistry:
    """Registry that selects an appropriate loader for a path."""

    def __init__(self) -> None:
        self._loaders: list[BaseLoader] = [
            MarkdownLoader(),
            TextLoader(),
            PDFLoader(),
            DocxLoader(),
            EmailLoader(),
            MboxLoader(),
        ]

    def register(self, loader: BaseLoader) -> None:
        self._loaders.append(loader)

    def for_path(self, path: Path) -> BaseLoader | None:
        for loader in self._loaders:
            if loader.can_load(path):
                return loader
        return None

    def load(self, path: Path) -> list[LoadedDocument]:
        loader = self.for_path(path)
        if loader is None:
            raise ValueError(f"No loader registered for suffix {path.suffix}")
        return loader.load(path)


def _split_front_matter(text: str) -> tuple[dict[str, object] | None, str]:
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            fm_raw = parts[1]
            body = parts[2]
            try:
                front_matter = yaml.safe_load(fm_raw) or {}
                if isinstance(front_matter, dict):
                    return front_matter, body
            except yaml.YAMLError:
                pass
    return None, text


def _markdown_to_text(text: str) -> str:
    tokens = _MD.parse(text)
    parts: list[str] = []
    for token in tokens:
        content = token.content.strip()
        if content:
            parts.append(content)
    return normalize("\n".join(parts) if parts else text)


def _detect_lang(text: str) -> str:
    if langid is None:
        return "en"
    try:
        lang, _ = langid.classify(text)
        return lang
    except Exception:
        return "en"


def _to_unix_ms(value: object | None) -> int | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return int(value.timestamp() * 1000)
    if isinstance(value, (int, float)):
        return int(float(value) * 1000)
    if isinstance(value, str):
        try:
            return int(datetime.fromisoformat(value).timestamp() * 1000)
        except ValueError:
            return None
    return None


def _file_time(ts: float) -> int:
    return int(ts * 1000)


def _email_to_document(path: Path, message: EmailMessage, index: int | None = None) -> LoadedDocument:
    subject = message.get("subject", "")
    from_addr = message.get("from")
    date_header = message.get("date")
    if message.is_multipart():
        parts: list[str] = []
        for part in message.walk():
            if part.get_content_type() == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    parts.append(payload.decode(part.get_content_charset() or "utf-8", errors="ignore"))
        body = "\n".join(parts)
    else:
        body_part = message.get_body(preferencelist=("plain",))
        if body_part:
            payload = body_part.get_content()
        else:
            payload = message.get_payload(decode=True) or b""
        body = payload if isinstance(payload, str) else payload.decode("utf-8", errors="ignore")

    normalized = normalize(body)
    metadata = {
        "path": str(path),
        "message_id": message.get("message-id"),
        "lang": _detect_lang(normalized) if normalized else "en",
    }
    size_bytes = len(normalized.encode("utf-8"))
    created_ts = None
    if date_header:
        try:
            parsed = email.utils.parsedate_to_datetime(date_header)
            created_ts = int(parsed.timestamp() * 1000)
        except Exception:
            created_ts = None
    external_id = f"{os.fspath(path)}#{index}" if index is not None else os.fspath(path)
    metadata["external_id"] = external_id
    return LoadedDocument(
        path=path,
        text=normalized,
        raw_bytes=None,
        metadata=metadata,
        mime="message/rfc822" if index is None else "message/rfc822+item",
        title=subject or path.stem,
        author=from_addr,
        created_ts=created_ts,
        modified_ts=_file_time(path.stat().st_mtime),
        size_bytes=size_bytes,
    )
